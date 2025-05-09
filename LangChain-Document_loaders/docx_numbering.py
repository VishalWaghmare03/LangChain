# from docx import Document
# from zipfile import ZipFile
# from lxml import etree

# def get_numbering_map(docx_path):
#     with ZipFile(docx_path) as docx_zip:
#         with docx_zip.open("word/numbering.xml") as numbering_file:
#             tree = etree.parse(numbering_file)
#             ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

#             num_map = {}
#             for num in tree.xpath("//w:num", namespaces=ns):
#                 num_id = num.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
#                 abstract_id = num.xpath(".//w:abstractNumId", namespaces=ns)[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
#                 num_map[num_id] = abstract_id
#             return num_map

# def extract_numbered_text(docx_path):
#     doc = Document(docx_path)
#     num_map = get_numbering_map(docx_path)
#     numbered_text = []

#     for para in doc.paragraphs:
#         numbering_props = para._p.pPr.numPr if para._p.pPr is not None else None
#         if numbering_props:
#             num_id = numbering_props.numId.val
#             ilvl = numbering_props.ilvl.val
#             prefix = f"{ilvl + 1}." * (ilvl + 1)  # crude multilevel prefix
#             text = f"{prefix} {para.text.strip()}"
#         else:
#             text = para.text.strip()

#         if text:
#             numbered_text.append(text)

#     return "\n".join(numbered_text)


# if __name__ == "__main__":
#     path = "/home/vishal/Documents/Documents/Non-Heading Document/4_4301.docx"
#     output = extract_numbered_text(path)
#     print(output)

from docx import Document

def extract_with_manual_numbering(docx_path):
    doc = Document(docx_path)
    numbered_text = []

    level_counters = {}  # Tracks counts per level

    for para in doc.paragraphs:
        ppr = para._p.pPr
        numbering_props = ppr.numPr if ppr is not None and ppr.numPr is not None else None

        if numbering_props is not None:
            ilvl = numbering_props.ilvl.val

            # Update counter
            if ilvl in level_counters:
                level_counters[ilvl] += 1
            else:
                level_counters[ilvl] = 1

            # Reset deeper levels
            for deeper in range(ilvl + 1, 10):
                level_counters.pop(deeper, None)

            # Compose numbering string like 1., 1.1., etc.
            prefix = ".".join(str(level_counters[i]) for i in sorted(level_counters) if i <= ilvl)
            text = f"{prefix}. {para.text.strip()}"
        else:
            text = para.text.strip()

        if text:
            numbered_text.append(text)

    return "\n".join(numbered_text)

if __name__ == "__main__":
    file_path = "/home/vishal/Documents/Documents/Non-Heading Document/Agreement for Sale.docx"  # <-- update this path
    extracted_text = extract_with_manual_numbering(file_path)

    print("\n✅ Extracted Numbered Text:\n")
    print(extracted_text)

    # Optional: Save to file
    with open("extracted_output.txt", "w", encoding="utf-8") as f:
        f.write(extracted_text)

